import os
import shutil
import time
import docx
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

class FileManipulator:
    def __init__(self, str_oldpath:str, str_newpath:str, max_file_dict:dict):
        self.str_oldpath = str_oldpath
        self.str_newpath = str_newpath
        self.max_file_dict = max_file_dict

    def cp_files(self):
        """
        将文件夹从 'str_oldpath' 复制到 'str_newpath'，文件夹中的文件名会根据各类别中文件名的最高数字索引进行更新。如果 'str_newpath' 目录已经存在，将创建一个带有时间戳的同名目录，并清空目录内容。
        """
        str_oldpath=self.str_oldpath
        str_newpath=self.str_newpath
        max_file_dict=self.max_file_dict
        # 检查源目录是否存在
        if not os.path.exists(str_oldpath):
            print("源目录不存在：", str_oldpath)
            return

        # 检查目标目录，如果已存在，则重命名为原名称_时间戳，并创建空的目标目录
        if os.path.exists(str_newpath):
            # 获取当前时间戳
            timestamp = int(time.time())
            # 构建新的目录名
            new_target_path = f"{str_newpath}_{timestamp}"
            # 重命名现有目录
            os.rename(str_newpath, new_target_path)
            # 创建新的空目录
            os.makedirs(str_newpath)
        else:
            # 如果目标目录不存在，直接创建
            os.makedirs(str_newpath)

        files = os.listdir(str_oldpath)

        for f_name in files:
            i_var1 = f_name.rfind("-")
            if i_var1 != -1:
                file_index = int(f_name[i_var1 + 1:])
                file_code = f_name[i_var1 + 1:]
                file_class = f_name[0:i_var1]

                last_index = int(max_file_dict.get(file_class, 0))

                if file_index > last_index:
                    max_file_dict[file_class] = file_code

        for key, value in max_file_dict.items():
            # 在新文件名中增加索引数字
            new_code = "{:04d}".format(int(value) + 1)
            # 源文件夹和目标文件夹路径
            source_folder = os.path.join(str_oldpath, f"{key}-{value}")
            target_folder = os.path.join(str_newpath, f"{key}-{new_code}")

            # 如果目标文件夹不存在，则复制源文件夹
            if not os.path.exists(target_folder):
                shutil.copytree(source_folder, target_folder)
                # 复制文件属性并设置时间戳
                shutil.copystat(source_folder, target_folder)
                os.utime(target_folder, (time.time(), time.time()))
            else:
                print(f"文件夹已存在于：{target_folder}")

        self.max_file_dict=max_file_dict

    def del_files(self):
        """
        删除给定目标路径下的文件

        参数:
            str_tarpath (str): 执行删除和重命名操作的目标路径。

        返回:
            None
        """
        str_tarpath=self.str_newpath

        def recursively_delete_contents(folder_path):
            """
            递归删除文件夹内的所有内容。

            参数:
                folder_path (str): 要删除内容的文件夹路径。

            返回:
                None
            """
            # 遍历文件夹中的所有子项
            for item_name in os.listdir(folder_path):
                item_path = os.path.join(folder_path, item_name)

                # 检查是否为文件或文件夹
                if os.path.isfile(item_path):
                    # 如果是文件，则删除
                    os.remove(item_path)
                elif os.path.isdir(item_path):
                    # 如果是文件夹，则递归删除该文件夹及内部所有内容
                    recursively_delete_contents(item_path)
                    # 删除空文件夹
                    os.rmdir(item_path)

        try:
            # 遍历目标路径下的所有子项：Analysis、Product...
            for folder_item in os.listdir(str_tarpath):
                folder_path = os.path.join(str_tarpath, folder_item)
                # 遍历文件夹内的所有子项：Data Pack、证据...
                for item_name in os.listdir(folder_path):
                    item_path = os.path.join(folder_path, item_name)

                    # 判断文件名以 ~$ 开头
                    if item_name.startswith('~$'):
                        # 删除临时文件
                        os.remove(item_path)
                    elif os.path.isdir(item_path):
                        # 如果是文件夹，则递归删除内部所有内容
                        recursively_delete_contents(item_path)
        except (FileNotFoundError, PermissionError, OSError) as e:
            print(f"发生错误：{e}")

    def ren_files(self):
        """
        对指定目录下文件进行重命名：0038-->0039
        """
        str_tarpath=self.str_newpath
        try:
            # 遍历目标路径下的所有子项：Analysis、Product...
            for folder_item in os.listdir(str_tarpath):
                folder_path = os.path.join(str_tarpath, folder_item)
                # 遍历文件夹内的所有子项：Data Pack、证据...
                for item_name in os.listdir(folder_path):
                    item_path = os.path.join(folder_path, item_name)

                    # 判断是文件还是文件夹
                    # 检查是否为文件，且文件名不以 ~$ 开头
                    if os.path.isfile(item_path) and not item_name.startswith('~$'):
                        # 查找括号的位置
                        start_bracket_index = item_name.find('(') if item_name.find('(') != -1 else item_name.find('（')
                        end_bracket_index = item_name.find(')') if item_name.find(')') != -1 else item_name.find('）')

                        # 检查是否找到了括号
                        if start_bracket_index != -1 and end_bracket_index != -1:
                            # 使用字符串切片和拼接来替换括号内的内容
                            new_file_name = item_name[:start_bracket_index + 1] + folder_item + item_name[end_bracket_index:]
                            # 定义新生成文件的路径
                            new_file_path = os.path.join(folder_path, new_file_name)

                            # 重命名文件
                            os.rename(item_path, new_file_path)
                        else:
                            print(f"文件名中没有找到括号，不进行重命名：{item_name}")
        except (FileNotFoundError, PermissionError, OSError) as e:
            print(f"发生错误：{e}")

    def edt_docx(self):
        str_tarpath=self.str_newpath
        def is_str_number(s):
            try:
                float(s)
                return True
            except ValueError:
                return False
        def run_paragraph(cells):
            for cell in cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.RED
            return
        # 获取当前脚本所在目录的绝对路径
        for head_file_name in os.listdir(str_tarpath):
            current_directory = os.path.join(str_tarpath,head_file_name)
            for file_name in os.listdir(current_directory):
                file_path=os.path.join(current_directory,file_name)
                if "REC-Q680003-A2" in file_name:
                    #如果是“REC-Q680003-A2-01  LIMS数据迁移表单”
                    # 加载现有的Word文档
                    doc = docx.Document(file_path)
                    tables = doc.tables

                    # 遍历文档中的所有段落
                    for tab in tables:
                        if "数据包名称" in tab.rows[0].cells[0].text:
                            # 修改表头
                            tab.rows[0].cells[1].text=head_file_name
                        else :
                            for t_row in tab.rows:
                                if is_str_number(t_row.cells[0].text):
                                    # 添加红色底纹
                                    run_paragraph(t_row.cells)

                    # 保存文档
                    doc.save(file_path)
                elif "REC-Q680003-A5" in file_name:
                    # 如果是“REC-Q680003-A5-01  LIMS主数据申请表”
                    # 加载现有的Word文档
                    doc = docx.Document(file_path)
                    tables = doc.tables
                    rows=tables[0].rows
                    rows_index=[2,3,5,7] # 添加红色底纹的行
                    rows[0].cells[2].text=head_file_name  # 修改表头
                    for r in rows_index:
                        run_paragraph(rows[r].cells)

                    
                    
                    for t_row in tables[2].rows:
                        if t_row.cells[0].text in head_file_name:
                            run_paragraph(t_row.cells)

                    # 保存文档
                    doc.save(file_path)

    def print_directory_tree(self, path:str, indent=0):
        """
        打印指定路径下的文件树。
        :param path: 要打印的目录路径
        :param indent: 当前缩进级别
        """
        for entry in os.listdir(path):
            entry_path = os.path.join(path, entry)
            if os.path.isfile(entry_path):
                print(' ' * indent + '|-- ' + entry)
            elif os.path.isdir(entry_path):
                print(' ' * indent + '|-- ' + entry)
                self.print_directory_tree(entry_path, indent + 4)

    def execute_operations(self):
        self.cp_files()
        self.del_files()
        self.ren_files()
        self.edt_docx()
        self.print_directory_tree(self.str_newpath)

# Example usage
str_oldpath = r'E:\WXQ_workspace\VScode\demo\oldf'
str_newpath = r'E:\WXQ_workspace\VScode\demo\newf'
max_file_dict = {}

fm = FileManipulator(str_oldpath, str_newpath, max_file_dict)
fm.execute_operations()
# fm.print_directory_tree(fm.str_newpath)