import re
import os
import json
import sys
import docx
from docx.enum.text import WD_COLOR_INDEX
import csv

if __name__ == "__main__":
    # 当直接运行此脚本时初始化配置
    config = FileUtils.load_config()  # 修正为使用类方法
    FileUtils.save_config(config)  # 修正为使用类方法

class FileUtils:
    """严格禁止实例化的工具类"""
    
    # 定义默认的head_list
    head_list = [
        'Analysis', 'CofA Step', 'Events', 'Format Calculation', 'LIMS Constant',
        'LIMS Users', 'Label Printer', 'Lists', 'Lot Template', 'Product',
        'Report Template', 'Stock', 'Subroutine', 'Suppliers', 'T PH AQL Sample Plan',
        'T PH Item Code', 'T PH Sample Plan', 'T_PH_Grade', 'T_PH_Spec Type', 'T_Report Text',
        'Table Master', 'Table Template', 'Units', 'User Dialog', 'vendor', 'Stage'
    ]

    @staticmethod
    def get_config_path():
        """获取配置文件路径（支持打包环境）"""
        if getattr(sys, 'frozen', False):
            # 打包环境 - 使用exe所在目录
            base_path = os.path.dirname(sys.executable)
        else:
            # 开发环境 - 使用脚本所在目录
            base_path = os.path.dirname(os.path.abspath(__file__))
        return os.path.join(base_path, 'config.json')

    @staticmethod
    def save_config(config):
        """保存配置到文件（支持打包环境）"""
        config_path = FileUtils.get_config_path()
        try:
            with open(config_path, 'w', encoding='utf-8') as f:
                json.dump(config, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"保存配置失败: {e}")

    @staticmethod
    def load_config():
        """从文件加载配置（支持打包环境）"""
        config = {
            'head_list': FileUtils.head_list,
            'default_old_path': '',
            'default_new_path': '',
        }
        config_path = FileUtils.get_config_path()
        
        try:
            if os.path.exists(config_path):
                with open(config_path, 'r', encoding='utf-8') as f:
                    loaded_config = json.load(f)
                    config.update(loaded_config)
            else:
                # 如果配置文件不存在，创建默认配置
                FileUtils.save_config(config)
        except Exception as e:
            print(f"加载配置失败: {e}")
        
        return config
        
    @staticmethod
    def recursively_delete_contents(folder_path):
        """
        递归删除文件夹内的所有内容。

        参数:
            folder_path (str): 要删除内容的文件夹路径。

        返回:
            None
        """
        # 遍历文件夹中的所有子项
        for item_name in os.listdir(folder_path):
            item_path = os.path.join(folder_path, item_name)

            # 检查是否为文件或文件夹
            if os.path.isfile(item_path):
                # 如果是文件，则删除
                os.remove(item_path)
            elif os.path.isdir(item_path):
                # 如果是文件夹，则递归删除该文件夹及内部所有内容
                FileUtils.recursively_delete_contents(item_path)
                # 删除空文件夹
                os.rmdir(item_path)

    @staticmethod
    def increment_filename_number(filename: str, start_sep: str = '', end_sep: str = '') -> str:
        """
        文件名数字递增方法
        
        参数:
            filename: 原始文件名
            start_sep: 起始定位字符串
            end_sep: 结束定位字符串
            
        返回:
            处理后的新文件名
        """
        if start_sep != '':
            # 找到起始定位字符串第一次出现的位置
            start_index = filename.find(start_sep)
            if start_index == -1:
                return filename  # 未找到起始定位字符串
            
            # 计算起始搜索位置（跳过起始字符串）
            start_search = start_index + len(start_sep)
        else :
            start_search = 0
            
        if end_sep != '':
            # 找到结束定位字符串最后一次出现的位置
            end_index = filename.rfind(end_sep, start_search)
            if end_index == -1:
                return filename  # 未找到结束定位字符串
        else:
            end_index = len(filename)
        
        # 提取目标范围内的内容
        target_str = filename[start_search:end_index]
        
        # 在目标字符串中查找第一个连续数字序列
        match = re.search(r'\d+', target_str)
        if not match:
            return filename  # 未找到数字序列
        
        # 提取数字部分和原始位数
        number_part = match.group()
        original_digits = len(number_part)
        start_pos = match.start()
        end_pos = match.end()
        
        # 转换为整数并+1
        try:
            number_value = int(number_part)
            new_number = number_value + 1
            new_number_str = str(new_number)
        except ValueError:
            return filename
        
        # 保留原始位数格式（前导零）
        if len(new_number_str) < original_digits:
            new_number_str = new_number_str.zfill(original_digits)
        
        # 构建新目标字符串（仅替换数字部分）
        new_target_str = (
            target_str[:start_pos] + 
            new_number_str + 
            target_str[end_pos:]
        )
        
        # 构建完整新文件名
        return (
            filename[:start_search] + 
            new_target_str + 
            filename[end_index:]
        )

    @staticmethod
    def is_str_number(s):
        try:
            float(s)
            return True
        except ValueError:
            return False

    @staticmethod
    def run_paragraph(cells):
        """加红色底纹"""
        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.highlight_color = WD_COLOR_INDEX.RED
        return

    @staticmethod
    def find_files_by_name(search_path, name_contains:str, extension=None, exclude_temp=True):
        """
        查找指定路径下文件名包含特定字符的文件
        
        参数:
        search_path (str): 要搜索的根目录路径
        name_contains (str): 文件名需要包含的字符串
        extension (str, optional): 文件后缀名（如'txt'或'.txt'），默认None表示不限后缀
        exclude_temp (bool, optional): 是否排除临时文件，默认True（排除）
        
        返回:
        list: 匹配文件的完整路径列表
        """
        matched_files = []
        
        # 标准化后缀格式（确保带点号）
        normalized_extension = None
        if extension is not None:
            # 处理带点号和不带点号的后缀
            normalized_extension = f".{extension.lstrip('.')}" if extension else ""
        
        # 遍历目录树
        for root, _, files in os.walk(search_path):
            for file in files:
                # 检查是否临时文件（如果需要排除）
                if exclude_temp:
                    # 排除常见临时文件模式
                    if file.startswith('~$') or file.startswith('.~') or \
                    file.endswith('.tmp') or file.endswith('.temp'):
                        continue
                
                # 检查文件名是否包含目标字符串
                if name_contains in file:
                    file_path = os.path.join(root, file)
                    
                    # 检查后缀条件
                    if normalized_extension is None:
                        matched_files.append(file_path)
                    else:
                        # 获取文件后缀并转换为小写比较（可选）
                        file_ext = os.path.splitext(file)[1]
                        if file_ext.lower() == normalized_extension.lower():
                            matched_files.append(file_path)
        
        return matched_files

    @staticmethod
    def edt_docx(doc_path: str, doc_name: str):
        file_path = os.path.join(doc_path, doc_name)

        # 使用类变量head_list
        if any(doc_name.startswith(head) for head in FileUtils.head_list):
            # 处理封面文件
            doc = docx.Document(file_path)
            tab = doc.tables[0]

            # 获取目标单元格
            target_cell = tab.rows[2].cells[0]

            # 保存原始格式的文本运行(runs)
            original_runs = target_cell.paragraphs[0].runs.copy()

            # 清除单元格内容
            for paragraph in target_cell.paragraphs:
                p = paragraph._element
                p.getparent().remove(p)

            # 添加新段落
            new_paragraph = target_cell.add_paragraph()

            # 添加新文本运行，并复制原始格式
            if original_runs:
                # 使用第一个原始运行的格式作为基准
                base_run = original_runs[0]
                new_run = new_paragraph.add_run(doc_name.replace(".docx", ""))

                # 复制字体格式
                new_run.font.name = base_run.font.name
                new_run.font.size = base_run.font.size
            else:
                # 如果没有原始运行，直接添加文本
                new_paragraph.add_run(doc_name.replace(".docx", ""))
            doc.save(file_path)
        elif "REC-Q680003-A2" in doc_name:
            # 如果是"REC-Q680003-A2-01  LIMS数据迁移表单"
            doc = docx.Document(file_path)
            tables = doc.tables

            for tab in tables:
                if tab.rows and "数据包名称" in tab.rows[0].cells[0].text:
                    # 修改表头
                    new_text = FileUtils.increment_filename_number(tab.rows[0].cells[1].text)
                    tab.rows[0].cells[1].text = new_text
                else:
                    for t_row in tab.rows:
                        if t_row.cells and FileUtils.is_str_number(t_row.cells[0].text):
                            # 添加红色底纹
                            FileUtils.run_paragraph(t_row.cells)
            doc.save(file_path)

        elif "REC-Q680003-A5" in doc_name:
            # 如果是"REC-Q680003-A5-01  LIMS主数据申请表"
            doc = docx.Document(file_path)
            tables = doc.tables
            rows = tables[0].rows
            rows_index = [2, 3, 5, 7]  # 添加红色底纹的行
            
            new_text = FileUtils.increment_filename_number(rows[0].cells[2].text)
            rows[0].cells[2].text = new_text
            
            for r in rows_index:
                if len(rows) > r:
                    FileUtils.run_paragraph(rows[r].cells)

            if len(tables) > 2:
                for t_row in tables[2].rows:
                    if t_row.cells and t_row.cells[0].text == new_text:
                        FileUtils.run_paragraph(t_row.cells)

            doc.save(file_path)

    @staticmethod
    def edit_A2_docx(current_directory: str, file_name: str, to_val_date: str, to_prod_date: str):
        """批量修改迁移日期"""
        file_path = os.path.join(current_directory, file_name)
        doc = docx.Document(file_path)
        tables = doc.tables

        for tab in tables:
            if tab.rows and "数据包名称" in tab.rows[0].cells[0].text:
                continue
                
            for t_row in tab.rows:
                if not t_row.cells:
                    continue
                    
                if FileUtils.is_str_number(t_row.cells[0].text):
                    if len(t_row.cells) < 5:
                        continue
                        
                    # 获取单元格引用
                    val_data_cell = t_row.cells[3]
                    prod_data_cell = t_row.cells[4]
                    
                    # 处理验证日期单元格
                    # 保存原始段落格式
                    val_alignment = None
                    if val_data_cell.paragraphs:
                        val_alignment = val_data_cell.paragraphs[0].alignment
                    
                    # 清除现有内容
                    val_paragraphs = list(val_data_cell.paragraphs)
                    for para in val_paragraphs:
                        p = para._element
                        p.getparent().remove(p)
                    
                    # 添加新内容并继承原始对齐方式
                    val_para = val_data_cell.add_paragraph()
                    val_run = val_para.add_run(to_val_date)
                    if val_alignment is not None:
                        val_para.alignment = val_alignment
                    
                    # 处理生产日期单元格
                    # 保存原始段落格式
                    prod_alignment = None
                    if prod_data_cell.paragraphs:
                        prod_alignment = prod_data_cell.paragraphs[0].alignment
                    
                    # 清除现有内容
                    prod_paragraphs = list(prod_data_cell.paragraphs)
                    for para in prod_paragraphs:
                        p = para._element
                        p.getparent().remove(p)
                    
                    # 添加新内容并继承原始对齐方式
                    prod_para = prod_data_cell.add_paragraph()
                    prod_run = prod_para.add_run(to_prod_date)
                    if prod_alignment is not None:
                        prod_para.alignment = prod_alignment
                    
                    # 设置字体格式（保留原有逻辑）
                    base_runs = None
                    if val_paragraphs and val_paragraphs[0].runs:
                        base_runs = val_paragraphs[0].runs
                    elif prod_paragraphs and prod_paragraphs[0].runs:
                        base_runs = prod_paragraphs[0].runs
                    
                    if base_runs:
                        base_run = base_runs[0]
                        val_run.font.name = base_run.font.name
                        val_run.font.size = base_run.font.size
                        prod_run.font.name = base_run.font.name
                        prod_run.font.size = base_run.font.size

        # 统一保存文档
        doc.save(file_path)

    @staticmethod
    def read_A2(path_a2: str) -> list:
        put_list=[]
        try:
            doc = docx.Document(path_a2)
            tbs = doc.tables
            
            # 查找包含"数据包名称"的表格
            pack_name_table_idx = None
            for tb_idx, tab in enumerate(tbs):
                try:
                    if "数据包名称" in tab.cell(0, 0).text.strip() :
                        pack_name_table_idx = tb_idx
                        break
                except IndexError:
                    continue
            
            if pack_name_table_idx is None:
                print(f"在文档 {os.path.basename(path_a2)} 中未找到'数据包名称'表格")
                return
            
            target_table = tbs[pack_name_table_idx +1]
            package_name = tbs[pack_name_table_idx].cell(0, 1).text.strip()
            
            # 处理目标表格中的行
            for r_idx,row in enumerate(target_table.rows):
                try:
                    cell0_text = row.cells[0].text.strip()
                    # 第一行为标题，跳过
                    if r_idx == 0:
                        continue 
                    # 如果第一列不是数字，停止处理               
                    elif not FileUtils.is_str_number(cell0_text):
                        break
                    
                    # 获取其他列的内容（注意索引从0开始）
                    record_name = row.cells[2].text.strip() if len(row.cells) > 2 else ""
                    to_val_date = row.cells[3].text.strip() if len(row.cells) > 3 else ""
                    to_prod_date = row.cells[4].text.strip() if len(row.cells) > 4 else ""
                    
                    put_list.append((package_name, record_name, to_val_date, to_prod_date))
                    
                except Exception as row_ex:
                    print(f"处理行时出错: {str(row_ex)}")
        
        except Exception as e:
            print(f"处理文档 {path_a2} 时出错: {str(e)}")
        return put_list

    @staticmethod
    def read_A5(path_a5: str) -> list:
        put_list = []
        tb2_list = []

        try:
            doc = docx.Document(path_a5)
            tbs = doc.tables
            
            # 添加表格存在性检查
            if len(tbs) < 3:
                print(f"文档表格不足: {path_a5}")
                return [[], []]
            
            pack_name = tbs[0].cell(0, 2).text.strip()  # 包名称
            justification = tbs[0].cell(5, 0).text.strip()  # 理由
            related_doc = tbs[0].cell(7, 0).text.strip()  # 相关文件

            tb1_list = (pack_name, justification, related_doc)

            # 处理目标表格中的行
            for row in tbs[2].rows:
                try:
                    # 添加单元格存在性检查
                    if not row.cells or len(row.cells) < 5:
                        continue
                    if row.cells[0].text.lower() != pack_name[0:len(pack_name)-5].lower():
                        continue
                    
                    record_name = row.cells[1].text.strip() 
                    oper_type = row.cells[2].text.strip() 
                    classification = row.cells[3].text.strip()
                    criticality  = row.cells[4].text.strip()
                    
                    
                    tb2_list.append((pack_name, record_name, oper_type, classification, criticality))
                    
                except Exception as row_ex:
                    print(f"处理行时出错: {str(row_ex)}")
            put_list = [tb1_list, tb2_list]
        except Exception as e:
            print(f"处理文档 {path_a5} 时出错: {str(e)}")
            put_list = [[], []]  # 确保返回两个空列表
        
        return put_list

    @staticmethod
    def write_to_csv(data: list, output_path: str,title:list=[]):
        """将数据写入CSV文件"""
        try:
            with open(output_path, 'w', newline='', encoding='utf-8-sig') as csvfile:
                writer = csv.writer(csvfile)
                
                if title:
                    # 写入表头
                    writer.writerow(title)
                # 写入所有数据行
                writer.writerows(data)
            print(f"成功写入CSV文件: {output_path}")
            return True
        except Exception as e:
            print(f"写入CSV失败: {str(e)}")
            return False