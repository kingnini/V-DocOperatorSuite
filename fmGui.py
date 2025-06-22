import os
import re
import shutil
import time
import docx
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, 
                            QLabel, QLineEdit, QPushButton, QTextEdit, QFileDialog, QGroupBox,
                            QFormLayout, QMessageBox)
from PyQt5.QtCore import Qt

class FileUtils:
    """严格禁止实例化的工具类"""

    @staticmethod
    def recursively_delete_contents(folder_path):
        """
        递归删除文件夹内的所有内容。

        参数:
            folder_path (str): 要删除内容的文件夹路径。

        返回:
            None
        """
        # 遍历文件夹中的所有子项
        for item_name in os.listdir(folder_path):
            item_path = os.path.join(folder_path, item_name)

            # 检查是否为文件或文件夹
            if os.path.isfile(item_path):
                # 如果是文件，则删除
                os.remove(item_path)
            elif os.path.isdir(item_path):
                # 如果是文件夹，则递归删除该文件夹及内部所有内容
                FileUtils.recursively_delete_contents(item_path)
                # 删除空文件夹
                os.rmdir(item_path)

    @staticmethod
    def increment_filename_number(filename: str, start_sep: str = '', end_sep: str = '') -> str:
        """
        文件名数字递增方法
        
        参数:
            filename: 原始文件名
            start_sep: 起始定位字符串
            end_sep: 结束定位字符串
            
        返回:
            处理后的新文件名
        """
        if start_sep != '':
            # 找到起始定位字符串第一次出现的位置
            start_index = filename.find(start_sep)
            if start_index == -1:
                return filename  # 未找到起始定位字符串
            
            # 计算起始搜索位置（跳过起始字符串）
            start_search = start_index + len(start_sep)
        else :
            start_search = 0
            
        if end_sep != '':
            # 找到结束定位字符串最后一次出现的位置
            end_index = filename.rfind(end_sep, start_search)
            if end_index == -1:
                return filename  # 未找到结束定位字符串
        else:
            end_index = len(filename)
        
        # 提取目标范围内的内容
        target_str = filename[start_search:end_index]
        
        # 在目标字符串中查找第一个连续数字序列
        match = re.search(r'\d+', target_str)
        if not match:
            return filename  # 未找到数字序列
        
        # 提取数字部分和原始位数
        number_part = match.group()
        original_digits = len(number_part)
        start_pos = match.start()
        end_pos = match.end()
        
        # 转换为整数并+1
        try:
            number_value = int(number_part)
            new_number = number_value + 1
            new_number_str = str(new_number)
        except ValueError:
            return filename
        
        # 保留原始位数格式（前导零）
        if len(new_number_str) < original_digits:
            new_number_str = new_number_str.zfill(original_digits)
        
        # 构建新目标字符串（仅替换数字部分）
        new_target_str = (
            target_str[:start_pos] + 
            new_number_str + 
            target_str[end_pos:]
        )
        
        # 构建完整新文件名
        return (
            filename[:start_search] + 
            new_target_str + 
            filename[end_index:]
        )

    @staticmethod
    def is_str_number(s):
        try:
            float(s)
            return True
        except ValueError:
            return False

    @staticmethod
    def run_paragraph(cells):
        """加红色底纹"""
        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.RED
        return

    @staticmethod
    def edt_docx(doc_path:str,doc_name:str):

        head_list=[
                'Analysis','CofA Step','Events','Format Calculation','LIMS Constant',
                'LIMS Users','Label Printer','Lists','Lot Template','Product',
                'Report Template','Stock','Subroutine','Suppliers','T PH AQL Sample Plan',
                'T PH Item Code','T PH Sample Plan','T_PH_Grade','T_PH_Spec Type','T_Report Text',
                'Table Master','Table Template','Units','User Dialog','vendor','Stage'
            ]
        file_path = os.path.join(doc_path, doc_name)

        if any(doc_name.startstartswith(head) for head in head_list)
            # 处理封面文件
            doc = docx.Document(file_path)
            tab = doc.tables[0]

            # 获取目标单元格
            target_cell = tab.rows[2].cells[0]

            # 保存原始格式的文本运行(runs)
            original_runs = target_cell.paragraphs[0].runs.copy()

            # 清除单元格内容
            for paragraph in target_cell.paragraphs:
                p = paragraph._element
                p.getparent().remove(p)

            # 添加新段落
            new_paragraph = target_cell.add_paragraph()

            # 添加新文本运行，并复制原始格式
            if original_runs:
                # 使用第一个原始运行的格式作为基准
                base_run = original_runs[0]
                new_run = new_paragraph.add_run(doc_name.replace(".docx", ""))

                # 复制字体格式
                new_run.font.name = base_run.font.name
                new_run.font.size = base_run.font.size
            else:
                # 如果没有原始运行，直接添加文本
                new_paragraph.add_run(doc_name.replace(".docx", ""))
            doc.save(file_path)
        elif "REC-Q680003-A2" in file_name:
            # 如果是“REC-Q680003-A2-01  LIMS数据迁移表单”
            # 加载现有的Word文档
            doc = docx.Document(file_path)
            tables = doc.tables

            # 遍历文档中的所有段落
            for tab in tables:
                if "数据包名称" in tab.rows[0].cells[0].text:
                    # 修改表头
                    tab.rows[0].cells[1].text = increment_filename_number(tab.rows[0].cells[1].text)
                else:
                    for t_row in tab.rows:
                        if FileUtils.is_str_number(t_row.cells[0].text):
                            # 添加红色底纹
                            FileUtils.run_paragraph(t_row.cells)

            # 保存文档
            doc.save(file_path)

        elif "REC-Q680003-A5" in file_name:
            # 如果是“REC-Q680003-A5-01  LIMS主数据申请表”
            # 加载现有的Word文档
            doc = docx.Document(file_path)
            tables = doc.tables
            rows = tables[0].rows
            rows_index = [2, 3, 5, 7]  # 添加红色底纹的行
            rows[0].cells[2].text = increment_filename_number(rows[0].cells[2].text)  # 修改表头
            for r in rows_index:
                FileUtils.run_paragraph(rows[r].cells)

            for t_row in tables[2].rows:
                if t_row.cells[0].text in rows[0].cells[2].text:
                    FileUtils.run_paragraph(t_row.cells)

            # 保存文档
            doc.save(file_path)

    @staticmethod
    def edit_A2_docx(current_directory:str,file_name:str,to_val_date:str,to_prod_date:str):
        """批量修改迁移日期"""
        file_path = os.path.join(doc_path, doc_name)

        if "REC-Q680003-A2" in file_name:
            # 如果是“REC-Q680003-A2-01  LIMS数据迁移表单”
            # 加载现有的Word文档
            doc = docx.Document(file_path)
            tables = doc.tables

            # 遍历文档中的所有段落
            for tab in tables:
                if "数据包名称" in tab.rows[0].cells[0].text:
                    # 修改表头
                    tab.rows[0].cells[1].text = increment_filename_number(tab.rows[0].cells[1].text)
                else:
                    for t_row in tab.rows:
                        if FileUtils.is_str_number(t_row.cells[0].text):
                            # 添加红色底纹
                            FileUtils.run_paragraph(t_row.cells)

            # 保存文档
            doc.save(file_path)
            
class FileManipulator:
    def __init__(self, str_oldpath: str, str_newpath: str, max_file_dict: dict, output_callback=None):
        self.str_oldpath = str_oldpath
        self.str_newpath = str_newpath
        self.max_file_dict = max_file_dict
        self.output_callback = output_callback

    def log(self, message):
        """记录日志信息，如果有回调函数则使用它，否则打印到控制台"""
        if self.output_callback:
            self.output_callback(message)
        else:
            print(message)

    def cp_files(self):
        """
        将文件夹从 'str_oldpath' 复制到 'str_newpath'，文件夹中的文件名会根据各类别中文件名的最高数字索引进行更新。
        如果 'str_newpath' 目录已经存在，会将目标文件复制移动到一个带有时间戳的同名目录，然后情况目标文件。
        """
        self.log("开始复制文件...")
        str_oldpath = self.str_oldpath
        str_newpath = self.str_newpath
        max_file_dict = self.max_file_dict
        
        # 检查源目录是否存在
        if not os.path.exists(str_oldpath):
            self.log(f"源目录不存在：{str_oldpath}")
            return False

        # 检查目标目录，如果已存在，则重命名为原名称_时间戳，并创建空的目标目录
        if os.path.exists(str_newpath):
            # 获取当前时间戳
            timestamp = int(time.time())
            # 构建新的目录名
            new_target_path = f"{str_newpath}_{timestamp}"
            # 重命名现有目录
            os.rename(str_newpath, new_target_path)
            self.log(f"目标目录已存在，已重命名为: {new_target_path}")
            # 创建新的空目录
            os.makedirs(str_newpath)
            self.log(f"已创建新目录: {str_newpath}")
        else:
            # 如果目标目录不存在，直接创建
            os.makedirs(str_newpath)
            self.log(f"已创建目标目录: {str_newpath}")

        files = os.listdir(str_oldpath)

        for f_name in files:
            i_var1 = f_name.rfind("-")  # -所在的位置
            # 新增封面文件判断逻辑处理
            if i_var1 != -1 and os.path.isdir(os.path.join(str_oldpath, f_name)):
                try:
                    file_index = int(f_name[i_var1 + 1:])
                    file_code = f_name[i_var1 + 1:]
                    file_class = f_name[0:i_var1]

                    last_index = int(max_file_dict.get(file_class, 0))

                    if file_index > last_index:
                        max_file_dict[file_class] = file_code
                        self.log(f"更新类别 '{file_class}' 的最大索引为: {file_code}")
                except ValueError:
                    self.log(f"文件名 '{f_name}' 的数字部分无效，跳过处理")

        for key, value in max_file_dict.items():
            # 在新文件名中增加索引数字
            try:
                new_code = "{:04d}".format(int(value) + 1)
            except ValueError:
                self.log(f"类别 '{key}' 的索引值 '{value}' 无效，跳过处理")
                continue

            # 源文件夹和目标文件夹路径
            source_folder = os.path.join(str_oldpath, f"{key}-{value}")
            target_folder = os.path.join(str_newpath, f"{key}-{new_code}")

            # 检查源文件夹是否存在
            if not os.path.exists(source_folder):
                self.log(f"源文件夹不存在：{source_folder}")
                continue

            # 如果目标文件夹不存在，则复制源文件夹
            if not os.path.exists(target_folder):
                try:
                    shutil.copytree(source_folder, target_folder)
                    # 复制文件属性并设置时间戳
                    shutil.copystat(source_folder, target_folder)
                    os.utime(target_folder, (time.time(), time.time()))
                    self.log(f"已复制: {source_folder} -> {target_folder}")
                except Exception as e:
                    self.log(f"复制文件夹时出错: {e}")
            else:
                self.log(f"文件夹已存在于：{target_folder}")

        self.max_file_dict = max_file_dict

        # 复制封面文件
        for f_name in files:
            source_file = os.path.join(str_oldpath, f_name)
            if not os.path.isdir(source_file):
                target_file = os.path.join(str_newpath, f_name)

                # 如果目标文件不存在，则复制
                if not os.path.exists(target_file):
                    try:
                        shutil.copy(source_file, target_file)
                        # 复制文件属性并设置时间戳
                        shutil.copystat(source_file, target_file)
                        os.utime(target_file, (time.time(), time.time()))
                        self.log(f"已复制文件: {f_name}")
                    except Exception as e:
                        self.log(f"复制文件时出错: {e}")
                else:
                    self.log(f"文件已存在于：{target_file}")

        self.log("文件复制完成")
        return True

    def del_files(self):
        """
        删除给定目标路径下的文件

        参数:
            str_tarpath (str): 执行删除操作的目标路径。

        返回:
            None
        """
        self.log("开始删除临时文件...")
        str_tarpath = self.str_newpath

        try:
            # 遍历目标路径下的所有子项：Analysis、Product...
            for folder_item in os.listdir(str_tarpath):
                folder_path = os.path.join(str_tarpath, folder_item)

                # 判断是否为文件夹
                if not os.path.isdir(folder_path):
                    continue  # 跳过非文件夹项

                # 遍历文件夹内的所有子项：Data Pack、证据...
                for item_name in os.listdir(folder_path):
                    item_path = os.path.join(folder_path, item_name)

                    # 判断文件名以 ~$ 开头
                    if item_name.startswith('~$'):
                        # 删除临时文件
                        try:
                            os.remove(item_path)
                            self.log(f"已删除临时文件: {item_name}")
                        except Exception as e:
                            self.log(f"删除临时文件时出错: {e}")
                    elif os.path.isdir(item_path):
                        # 如果是文件夹，则递归删除内部所有内容
                        try:
                            FileUtils.recursively_delete_contents(item_path)
                            self.log(f"已删除文件夹: {item_path}")
                        except Exception as e:
                            self.log(f"删除文件夹时出错: {e}")
        except (FileNotFoundError, PermissionError, OSError) as e:
            self.log(f"发生错误：{e}")
            return False

        self.log("临时文件删除完成")
        return True

    def ren_files(self):
        """
        对指定目录下文件进行重命名：0038-->0039
        """
        self.log("开始重命名文件...")
        str_tarpath = self.str_newpath
        try:
            # 遍历目标路径下的所有子项：Analysis、Product...
            for folder_item in os.listdir(str_tarpath):
                folder_path = os.path.join(str_tarpath, folder_item)

                # 判断是否为文件夹
                if not os.path.isdir(folder_path):
                    # 不是文件夹  则是封面文件
                    newFileName = FileUtils.increment_filename_number(folder_item)
                    new_file_path = os.path.join(str_tarpath, newFileName)
                    try:
                        os.rename(folder_path, new_file_path)  # 使用完整路径重命名
                        self.log(f"已重命名文件: {folder_item} -> {newFileName}")
                    except Exception as e:
                        self.log(f"重命名文件时出错: {e}")
                    continue  # 跳过非文件夹项

                # 遍历文件夹内的所有子项：Data Pack、证据...
                for item_name in os.listdir(folder_path):
                    item_path = os.path.join(folder_path, item_name)

                    # 检查是否为文件，且文件名不以 ~$ 开头
                    if os.path.isfile(item_path) and not item_name.startswith('~$'):
                        # 查找括号的位置
                        start_bracket_index = item_name.find('(') if item_name.find('(') != -1 else item_name.find('（')
                        end_bracket_index = item_name.find(')') if item_name.find(')') != -1 else item_name.find('）')

                        # 检查是否找到了括号
                        if start_bracket_index != -1 and end_bracket_index != -1:
                            # 使用字符串切片和拼接来替换括号内的内容
                            new_file_name = item_name[:start_bracket_index + 1] + folder_item + item_name[end_bracket_index:]
                            # 定义新生成文件的路径
                            new_file_path = os.path.join(folder_path, new_file_name)

                            # 重命名文件
                            try:
                                os.rename(item_path, new_file_path)
                                self.log(f"已重命名: {item_name} -> {new_file_name}")
                            except Exception as e:
                                self.log(f"重命名文件时出错: {e}")
                        else:
                            self.log(f"文件名中没有找到括号，不进行重命名：{item_name}")
        except (FileNotFoundError, PermissionError, OSError) as e:
            self.log(f"发生错误：{e}")
            return False

        self.log("文件重命名完成")
        return True

    def edt_docx(self):
        self.log("开始编辑Word文档...")
        str_tarpath = self.str_newpath

        # 获取当前脚本所在目录的绝对路径
        for head_file_name in os.listdir(str_tarpath):
            current_directory = os.path.join(str_tarpath, head_file_name)

            # 判断是否为文件夹
            if not os.path.isdir(current_directory):
                # 不是文件夹  则是封面文件
                try:
                    FileUtils.edt_docx(str_tarpath,head_file_name)
                    self.log(f"已编辑封面: {head_file_name}")
                except Exception as e:
                    self.log(f"编辑封面文件时出错: {e}")
                continue  # 跳过后续

            for file_name in os.listdir(current_directory):
                file_path = os.path.join(current_directory, file_name)
                try:
                    if "REC-Q680003-A2" in file_name:
                        # 如果是“REC-Q680003-A2-01  LIMS数据迁移表单”
                       FileUtils.edt_docx(current_directory,file_name)
                        self.log(f"已编辑迁移表: {file_name}")

                    elif "REC-Q680003-A5" in file_name:
                        # 如果是“REC-Q680003-A5-01  LIMS主数据申请表”
                        FileUtils.edt_docx(current_directory,file_name)
                        self.log(f"已编辑申请表: {file_name}")
                except Exception as e:
                    self.log(f"编辑Word文档时出错: {e}")

        self.log("Word文档编辑完成")
        return True
    
    def edt_A2_docx(self,to_val_date:str,to_prod_date:str =''):
        self.log("批量设置验证日期和迁移日期...")
        str_tarpath = self.str_newpath

        if to_val_date=='':
            return False
        else if to_val_date=='':
            to_val_date=to_val_date
        
        # 获取当前脚本所在目录的绝对路径
        for head_file_name in os.listdir(str_tarpath):
            current_directory = os.path.join(str_tarpath, head_file_name)

            # 判断是否为文件夹
            if not os.path.isdir(current_directory):
                continue  # 跳过后续

            for file_name in os.listdir(current_directory):
                file_path = os.path.join(current_directory, file_name)
                try:
                    if "REC-Q680003-A2" in file_name:
                        # 如果是“REC-Q680003-A2-01  LIMS数据迁移表单”
                       FileUtils.edit_A2_docx(current_directory,file_name,to_val_date,to_prod_date)
                        self.log(f"已修改: {file_name}")
                except Exception as e:
                    self.log(f"编辑Word文档时出错: {e}")

        self.log("Word文档编辑完成")
        return True

    def get_directory_tree(self, path: str, indent=0):
        """
        获取指定路径下的文件树结构。
        :param path: 要打印的目录路径
        :param indent: 当前缩进级别
        :return: 目录树字符串
        """
        tree = ""
        try:
            for entry in os.listdir(path):
                entry_path = os.path.join(path, entry)
                if os.path.isfile(entry_path):
                    tree += ' ' * indent + '|-- ' + entry + '\n'
                elif os.path.isdir(entry_path):
                    tree += ' ' * indent + '|-- ' + entry + '\n'
                    tree += self.get_directory_tree(entry_path, indent + 4)
        except Exception as e:
            self.log(f"获取目录树时出错: {e}")
        return tree

    def execute_operations(self):
        self.log("=" * 50)
        self.log("开始执行文件操作流程")
        self.log("=" * 50)
        
        if not self.cp_files():
            self.log("文件复制失败，中止操作")
            return False
            
        if not self.del_files():
            self.log("文件删除失败，中止操作")
            return False
            
        if not self.ren_files():
            self.log("文件重命名失败，中止操作")
            return False
            
        if not self.edt_docx():
            self.log("文件编辑失败，中止操作")
            return False
            
        self.log("=" * 50)
        self.log("所有操作成功完成")
        self.log("=" * 50)
        return True


class FileManagerApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("文件操作工具")
        self.setGeometry(100, 100, 800, 600)
        
        # 主控件和布局
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)
        
        # 路径输入区域
        path_group = QGroupBox("路径设置")
        path_layout = QFormLayout()
        
        self.old_path_edit = QLineEdit()
        self.old_path_edit.setPlaceholderText("选择源文件夹路径...")
        self.old_path_button = QPushButton("浏览...")
        self.old_path_button.clicked.connect(self.browse_old_path)
        
        old_path_layout = QHBoxLayout()
        old_path_layout.addWidget(self.old_path_edit)
        old_path_layout.addWidget(self.old_path_button)
        
        self.new_path_edit = QLineEdit()
        self.new_path_edit.setPlaceholderText("选择目标文件夹路径...")
        self.new_path_button = QPushButton("浏览...")
        self.new_path_button.clicked.connect(self.browse_new_path)
        
        new_path_layout = QHBoxLayout()
        new_path_layout.addWidget(self.new_path_edit)
        new_path_layout.addWidget(self.new_path_button)
        
        path_layout.addRow("源文件夹:", old_path_layout)
        path_layout.addRow("目标文件夹:", new_path_layout)
        path_group.setLayout(path_layout)
        
        # 按钮区域
        button_layout = QHBoxLayout()
        self.execute_button = QPushButton("执行文件操作")
        self.execute_button.clicked.connect(self.execute_operations)
        self.execute_button.setStyleSheet("background-color: #4CAF50; color: white; font-weight: bold;")
        
        self.tree_button = QPushButton("显示目录结构")
        self.tree_button.clicked.connect(self.show_directory_tree)
        self.tree_button.setStyleSheet("background-color: #2196F3; color: white; font-weight: bold;")
        
        button_layout.addWidget(self.execute_button)
        button_layout.addWidget(self.tree_button)
        
        # 日志输出区域
        log_group = QGroupBox("操作日志")
        log_layout = QVBoxLayout()
        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        log_layout.addWidget(self.log_text)
        log_group.setLayout(log_layout)
        
        # 添加组件到主布局
        main_layout.addWidget(path_group)
        main_layout.addLayout(button_layout)
        main_layout.addWidget(log_group, 1)
        
        # 初始化文件管理器
        self.file_manipulator = None

    def browse_old_path(self):
        path = QFileDialog.getExistingDirectory(self, "选择源文件夹")
        if path:
            self.old_path_edit.setText(path)

    def browse_new_path(self):
        path = QFileDialog.getExistingDirectory(self, "选择目标文件夹")
        if path:
            self.new_path_edit.setText(path)

    def log_message(self, message):
        """将消息添加到日志区域"""
        self.log_text.append(message)
        # 自动滚动到底部
        self.log_text.verticalScrollBar().setValue(
            self.log_text.verticalScrollBar().maximum()
        )

    def execute_operations(self):
        """执行文件操作"""
        old_path = self.old_path_edit.text().strip()
        new_path = self.new_path_edit.text().strip()
        
        if not old_path or not new_path:
            QMessageBox.warning(self, "路径错误", "请选择源文件夹和目标文件夹路径！")
            return
            
        if not os.path.exists(old_path):
            QMessageBox.warning(self, "路径错误", "源文件夹路径不存在！")
            return
            
        # 清空日志
        self.log_text.clear()
        
        # 创建文件操作器
        self.file_manipulator = FileManipulator(
            old_path, new_path, {}, self.log_message
        )
        
        # 执行操作
        self.log_message("开始文件操作流程...")
        success = self.file_manipulator.execute_operations()
        
        if success:
            self.log_message("\n✅ 所有操作成功完成！")
        else:
            self.log_message("\n❌ 操作过程中出现错误！")

    def show_directory_tree(self):
        """显示目录结构"""
        path = self.new_path_edit.text().strip()
        if not path:
            QMessageBox.warning(self, "路径错误", "请选择目标文件夹路径！")
            return
            
        if not os.path.exists(path):
            QMessageBox.warning(self, "路径错误", "目标文件夹路径不存在！")
            return
            
        # 清空日志
        self.log_text.clear()
        
        if not self.file_manipulator:
            self.file_manipulator = FileManipulator(
                "", path, {}, self.log_message
            )
            
        self.log_message("目录结构:\n")
        tree = self.file_manipulator.get_directory_tree(path)
        self.log_message(tree)


if __name__ == "__main__":
    app = QApplication([])
    window = FileManagerApp()
    window.show()
    app.exec_()