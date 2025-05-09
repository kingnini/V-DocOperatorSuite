import os
import shutil
import time
from pathlib import Path
from typing import Dict, Optional, Callable
import docx
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_COLOR_INDEX


class FileOperator:
    """处理文件系统操作的基础类"""
    
    def __init__(self, log_callback: Optional[Callable] = print):
        self.log = log_callback if log_callback else lambda _: None

    @staticmethod
    def _handle_existing_directory(target_path: Path) -> None:
        """处理已存在的目标目录"""
        if target_path.exists():
            timestamp = int(time.time())
            backup_path = target_path.parent / f"{target_path.name}_{timestamp}"
            shutil.move(str(target_path), str(backup_path))
        target_path.mkdir(parents=True, exist_ok=True)

    def _remove_contents(self, path: Path) -> None:
        """递归删除目录内容"""
        try:
            for item in path.iterdir():
                if item.is_file():
                    item.unlink()
                elif item.is_dir():
                    self._remove_contents(item)
                    item.rmdir()
        except Exception as e:
            self.log(f"删除错误 {path}: {str(e)}")
            raise


class DocxEditor:
    """处理DOCX文件操作"""
    
    HIGHLIGHT_COLOR = WD_COLOR_INDEX.RED
    
    @classmethod
    def _apply_highlight(cls, cell) -> None:
        """为单元格内容应用高亮"""
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.highlight_color = cls.HIGHLIGHT_COLOR

    @classmethod
    def process_a2_document(cls, doc_path: Path, header: str) -> None:
        """处理REC-Q680003-A2类型文档"""
        doc = docx.Document(doc_path)
        table = doc.tables[0]
        
        # 更新表头
        table.rows[0].cells[1].text = header
        
        # 处理数值单元格
        for row in table.rows:
            if row.cells[0].text.strip().isdigit():
                for cell in row.cells:
                    cls._apply_highlight(cell)
        
        doc.save(doc_path)

    @classmethod
    def process_a5_document(cls, doc_path: Path, header: str) -> None:
        """处理REC-Q680003-A5类型文档"""
        doc = docx.Document(doc_path)
        main_table = doc.tables[0]
        data_table = doc.tables[2]

        # 更新主表头
        main_table.rows[0].cells[2].text = header
        
        # 高亮指定行
        highlight_rows = [2, 3, 5, 7]
        for idx in highlight_rows:
            cls._apply_highlight(main_table.rows[idx].cells[0])

        # 处理数据表
        for row in data_table.rows:
            if row.cells[0].text == header:
                for cell in row.cells:
                    cls._apply_highlight(cell)
        
        doc.save(doc_path)


class FileManipulator(FileOperator):
    def __init__(
        self,
        source_path: str,
        target_path: str,
        max_index: Dict[str, int],
        log_callback: Optional[Callable] = print
    ):
        super().__init__(log_callback)
        self.source = Path(source_path)
        self.target = Path(target_path)
        self.max_index = max_index
        self._validate_paths()

    def _validate_paths(self) -> None:
        """验证路径有效性"""
        if not self.source.exists():
            raise FileNotFoundError(f"源目录不存在: {self.source}")
        if not self.source.is_dir():
            raise ValueError(f"源路径不是目录: {self.source}")

    def copy_files(self) -> None:
        """执行文件复制操作"""
        self._handle_existing_directory(self.target)
        self._update_max_index()
        self._copy_folders()

    def _update_max_index(self) -> None:
        """更新最大索引字典"""
        for item in self.source.iterdir():
            if "-" in item.name and item.is_dir():
                prefix, _, number = item.name.partition("-")
                current_num = int(number)
                if current_num > self.max_index.get(prefix, 0):
                    self.max_index[prefix] = current_num

    def _copy_folders(self) -> None:
        """执行实际的文件夹复制"""
        for prefix, last_num in self.max_index.items():
            new_num = f"{last_num + 1:04d}"
            # 修正源路径格式化为4位数字
            source = self.source / f"{prefix}-{last_num:04d}"  # 关键修改点
            dest = self.target / f"{prefix}-{new_num}"

            if not dest.exists():
                shutil.copytree(source, dest)
                shutil.copystat(source, dest)
                dest.touch()  # 更新访问时间
            else:
                self.log(f"文件夹已存在: {dest}")

    def clean_target(self) -> None:
        """清理目标目录"""
        try:
            for entry in self.target.iterdir():
                if entry.is_dir():
                    self._clean_directory(entry)
        except Exception as e:
            self.log(f"清理错误: {str(e)}")
            raise

    # def _clean_directory(self, path: Path) -> None:
    #     """清理单个目录内容"""
    #     for item in path.iterdir():
    #         if item.name.startswith("~$"):  # 临时文件
    #             item.unlink()
    #         elif item.is_dir():
    #             self._remove_contents(item)
    #             item.rmdir()
    def _clean_directory(self, path: Path) -> None:
        """清理单个目录内容（修正：保留目录结构）"""
        for item in path.iterdir():
            if item.name.startswith("~$"):  # 临时文件
                item.unlink()
            elif item.is_dir():
                # 仅删除子目录内容但保留目录本身
                self._remove_contents(item)
                try:
                    item.rmdir()  # 尝试删除空目录
                except OSError:
                    pass  # 允许目录非空时保留

    def rename_files(self) -> None:
        """执行文件重命名操作"""
        for category_dir in self.target.iterdir():
            if not category_dir.is_dir():
                continue
                
            for item in category_dir.iterdir():
                if item.is_file() and not item.name.startswith("~$"):
                    self._process_file_rename(item, category_dir.name)

    def _process_file_rename(self, file_path: Path, category: str) -> None:
        """处理单个文件的重命名"""
        name = file_path.stem
        for bracket in ("(", "（"):
            if bracket in name:
                left_idx = name.find(bracket)
                right_idx = name.find(")" if bracket == "(" else "）")
                if right_idx != -1:
                    # 修正括号闭合问题
                    new_name = f"{name[:left_idx+1]}{category}{name[right_idx:]}"
                    file_path.rename(file_path.with_name(new_name + file_path.suffix))
                    return
        self.log(f"未找到括号: {file_path.name}")

    def process_documents(self) -> None:
        """处理所有文档编辑"""
        for category_dir in self.target.iterdir():
            if not category_dir.is_dir():
                continue
                
            for doc_file in category_dir.glob("*.docx"):
                if "REC-Q680003-A2" in doc_file.name:
                    DocxEditor.process_a2_document(doc_file, category_dir.name)
                elif "REC-Q680003-A5" in doc_file.name:
                    DocxEditor.process_a5_document(doc_file, category_dir.name)

    def execute_workflow(self) -> None:
        """执行完整工作流程"""
        operations = [
            self.copy_files,
            self.clean_target,
            self.rename_files,
            self.process_documents
        ]
        
        for op in operations:
            try:
                op()
            except Exception as e:
                self.log(f"操作失败: {str(e)}")
                raise

        self.print_directory_tree()

    def print_directory_tree(self, path: Optional[Path] = None, indent: int = 0) -> None:
        """打印目录结构"""
        path = path or self.target
        prefix = " " * indent + "|-- "
        
        self.log(prefix + path.name)
        if path.is_dir():
            for item in path.iterdir():
                self.print_directory_tree(item, indent + 4)


# 使用示例
if __name__ == "__main__":
    config = {
        "source": "E:/WXQ_workspace/VScode/demo/oldf",
        "target": "E:/WXQ_workspace/VScode/demo/newf",
        "max_index": {}
    }

    try:
        fm = FileManipulator(
            source_path=config["source"],
            target_path=config["target"],
            max_index=config["max_index"]
        )
        fm.execute_workflow()
    except Exception as e:
        print(f"流程执行失败: {str(e)}")
